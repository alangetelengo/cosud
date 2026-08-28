"""Génère docs/Programme-reprise-registre-aout-209-courriers.docx"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOCS_DIR = os.path.dirname(__file__)
OUTPUT = os.path.join(DOCS_DIR, "Programme-reprise-registre-aout-209-courriers.docx")

GREEN = RGBColor(0x00, 0xA0, 0x55)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)


def set_run_font(run, *, size=11, bold=False, color=DARK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_heading_custom(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=GREEN)
    else:
        set_run_font(run, size=13 if level == 2 else 11, bold=True, color=DARK)
    return p


def add_para(doc: Document, text: str, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0])
    else:
        set_run_font(p.add_run(text))


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0])
    else:
        set_run_font(p.add_run(text))


def set_cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "06A269")
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F1F5F9")
    doc.add_paragraph()
    return table


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(64)
    section.right_margin = Pt(64)

    # En-tête
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("COSUD — ACSI")
    set_run_font(run, size=14, bold=True, color=GREEN)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Programme de reprise — Registre d’arrivée août\n"
        "209 courriers (tous types) — Commission DG"
    )
    set_run_font(run, size=16, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Période d’exécution : samedi, dimanche et lundi\n"
        "Périmètre : factures, MAD, autres courriers — puis dettes & moratoires"
    )
    set_run_font(run, size=10, italic=True, color=GRAY)

    # 1. Objectif
    add_heading_custom(doc, "1. Objectif", 1)
    add_para(
        doc,
        "Saisir dans COSUD l’ensemble des 209 courriers d’arrivée du registre d’août "
        "(factures, MAD, lettres, demandes et autres), sous forme de commission constituée "
        "par le Directeur Général, afin de démarrer l’exploitation courante avec une base à jour.",
    )
    add_bullet(doc, "Priorité 1 : enregistrer les 209 courriers (registre complet).")
    add_bullet(doc, "Priorité 2 : consolider dettes fournisseurs / prestataires.")
    add_bullet(doc, "Priorité 3 : moratoires — uniquement après les factures saisies.")

    # 2. Cadence
    add_heading_custom(doc, "2. Cadence sur 3 jours", 1)
    add_para(
        doc,
        "Volume : 209 courriers ÷ 3 jours ≈ 70 courriers / jour pour la commission. "
        "Avec 5 personnes : ≈ 14 saisies / personne / jour.",
    )
    add_table(
        doc,
        ["Jour", "Objectif collectif", "Focus"],
        [
            ["Samedi", "~70 courriers", "Briefing, tri des piles, démarrage intensif, référentiel"],
            ["Dimanche", "~70 courriers", "Volume — viser la fin des factures et MAD"],
            ["Lundi", "~69 + contrôle", "Finir le reste, contrôle qualité, dettes si temps"],
        ],
    )

    # 3. Composition
    add_heading_custom(doc, "3. Composition de la commission", 1)
    add_table(
        doc,
        ["Personne", "Rôle métier", "Focus saisie reprise"],
        [
            [
                "Mme Taty",
                "Secrétaire DG — responsable factures / prestataires",
                "Pile FACTURES + référentiel fournisseurs / prestataires",
            ],
            [
                "Mme Eleni",
                "Secrétaire DG — responsable suivi des dépenses",
                "Pile MAD (+ appui factures si besoin)",
            ],
            [
                "Mme Mireille",
                "Secrétaire DG — autres courriers",
                "Pile AUTRES (hors MAD et facture prestataire)",
            ],
            [
                "Mme Lydia",
                "Trésorière auprès de l’Agent comptable",
                "Renfort FACTURES / MAD (volume)",
            ],
            [
                "Mme Lebanitou",
                "Agent comptable",
                "Renfort FACTURES / MAD (volume)",
            ],
        ],
    )
    add_para(
        doc,
        "Accompagnement technique COSUD : coordination, pointage Excel, déblocage "
        "(doublons, scans, fiches manquantes) — pas une 6ᵉ file de saisie sauf besoin.",
        italic=True,
        size=10,
    )

    # 4. Tri
    add_heading_custom(doc, "4. Tri du registre en 3 piles (samedi matin)", 1)
    add_para(
        doc,
        "Avant toute saisie, découper physiquement ou sur Excel les 209 courriers :",
    )
    add_table(
        doc,
        ["Pile", "Contenu", "Pilote", "Renfort"],
        [
            ["1 — Factures", "Type facture / prestataire", "Mme Taty", "Mme Lydia, Mme Lebanitou"],
            ["2 — MAD", "MAD / états de besoins", "Mme Eleni", "Mme Lydia / Lebanitou si dispo"],
            ["3 — Autres", "Lettres, demandes, etc.", "Mme Mireille", "Appui ½ journée si pile très grosse"],
        ],
    )
    add_para(
        doc,
        "Règle anti-doublon : une plage de n° de registre (ou une pile) = une personne. "
        "Personne ne saisit la plage d’un autre.",
        bold=True,
    )

    # 5. Référentiel
    add_heading_custom(doc, "5. Référentiel fournisseurs / prestataires (critique)", 1)
    add_para(
        doc,
        "Depuis la mise en place du référentiel COSUD, une facture ne peut être enregistrée "
        "sans choisir une fiche fournisseur / prestataire / partenaire active.",
    )
    add_numbered(
        doc,
        "Dès qu’une facture concerne un fournisseur absent du référentiel : file d’attente chez Mme Taty (1–2 min).",
    )
    add_numbered(doc, "Mme Taty crée la fiche (nom, type, téléphone, e-mail si connu).")
    add_numbered(doc, "La secrétaire reprend immédiatement la saisie du courrier.")
    add_para(
        doc,
        "Sans ce poste « référentiel », toute la pile Factures s’arrête. C’est le principal goulot des 3 jours.",
        bold=True,
    )

    # 6. Mode opératoire
    add_heading_custom(doc, "6. Mode opératoire par courrier", 1)
    add_numbered(doc, "Identifier le type (facture / MAD / autre).")
    add_numbered(
        doc,
        "Saisir le courrier arrivée dans COSUD : n° registre, dates, objet, expéditeur, scan(s).",
    )
    add_numbered(
        doc,
        "Si FACTURE : choisir le fournisseur dans le référentiel + montant + téléphone + service demandeur.",
    )
    add_numbered(doc, "Si MAD : renseigner les champs demandés (émetteur, service, montant le cas échéant).")
    add_numbered(doc, "Cocher sur la checklist : « saisi ».")
    add_para(
        doc,
        "Pendant la saisie de masse : ne pas lancer les moratoires ni forcer tout le circuit paiement "
        "(sauf consigne contraire du DG). Objectif = registre complet dans COSUD.",
        italic=True,
    )

    # 7. Planning détaillé
    add_heading_custom(doc, "7. Planning détaillé", 1)

    add_heading_custom(doc, "7.1 Samedi — démarrage + volume", 2)
    add_bullet(doc, "20–30 min : briefing commission + attribution des piles / plages de n°.")
    add_bullet(doc, "Tri registre août → 3 piles (Factures / MAD / Autres).")
    add_bullet(doc, "Ouverture checklist Excel partagée (voir § 8).")
    add_bullet(doc, "Saisie intensive (~70). Mme Taty gère le référentiel au fil de l’eau.")
    add_bullet(doc, "Fin de journée (15 min) : pointage — saisis / restants / bloqués.")

    add_heading_custom(doc, "7.2 Dimanche — volume", 2)
    add_bullet(doc, "Même organisation (~70).")
    add_bullet(doc, "Objectif : terminer si possible les piles Factures et MAD.")
    add_bullet(doc, "Mme Taty : rattrapage fiches + contrôle rapide des factures déjà saisies.")

    add_heading_custom(doc, "7.3 Lundi — finition + qualité", 2)
    add_bullet(doc, "Finir le reste (~69), surtout pile Autres.")
    add_bullet(doc, "1–2 h de contrôle croisé : n° = papier, scan présent, montant facture, fournisseur.")
    add_bullet(
        doc,
        "Si le DG exige dettes / moratoires le lundi : Mme Taty + Mme Eleni après registre factures OK ; "
        "Mme Lydia / Mme Lebanitou en appui lecture montants. Pas de moratoire avant dettes cohérentes.",
    )

    # 8. Checklist
    add_heading_custom(doc, "8. Checklist Excel (colonnes recommandées)", 1)
    add_para(doc, "Une ligne = un courrier du registre papier d’août.")
    add_table(
        doc,
        ["Colonne", "Utilité"],
        [
            ["N° registre", "Identifiant papier / COSUD"],
            ["Date réception", "Contrôle"],
            ["Type", "facture / MAD / autre"],
            ["Expéditeur / fournisseur", "Libellé"],
            ["Objet", "Résumé"],
            ["Montant", "Obligatoire si facture"],
            ["Téléphone", "Obligatoire facture / demande (SMS)"],
            ["Payé ? / Reliquat ?", "Aide dettes après saisie"],
            ["Scan OK", "Oui / Non"],
            ["Secrétaire", "Qui saisit"],
            ["Statut COSUD", "à saisir → saisi → contrôlé"],
            ["Bloquant", "ex. fournisseur absent, scan manquant"],
        ],
    )

    # 9. Dettes et moratoires
    add_heading_custom(doc, "9. Dettes et moratoires (après le registre)", 1)
    add_para(
        doc,
        "Les modules dettes et moratoires s’appuient sur les factures saisies. Ordre obligatoire :",
    )
    add_numbered(doc, "Terminer la saisie des factures d’août.")
    add_numbered(doc, "Vérifier le cumul des dettes (Mme Taty — suivi factures / dettes).")
    add_numbered(
        doc,
        "Créer les plans de moratoire uniquement pour les fournisseurs avec dette > 0 "
        "(instruction DG, échéances) — Mme Taty / circuit DG selon procédures COSUD.",
    )
    add_para(
        doc,
        "Ne pas créer de moratoire pendant que les 209 ne sont pas encore dans COSUD : "
        "risque de dette fausse.",
        bold=True,
    )

    # 10. Message
    add_heading_custom(doc, "10. Message à lire à la commission (samedi)", 1)
    add_para(
        doc,
        "« Nous avons 3 jours pour saisir les 209 arrivées d’août, tous types. "
        "Chacun a sa pile : Taty = factures + référentiel ; Eleni = MAD ; "
        "Mireille = les autres courriers ; Lydia et Lebanitou = renfort factures/MAD. "
        "Scan obligatoire. Pour une facture, le fournisseur se choisit dans le référentiel, "
        "pas en saisie libre. On ne crée pas les moratoires pendant la saisie : "
        "d’abord le registre complet dans COSUD. »",
        italic=True,
    )

    # 11. Points d'attention
    add_heading_custom(doc, "11. Points d’attention", 1)
    add_bullet(doc, "Doublons de n° registre : COSUD refuse — respecter les plages.")
    add_bullet(doc, "Facture sans fiche référentiel : bloqué jusqu’à création par Mme Taty.")
    add_bullet(doc, "Téléphone obligatoire sur facture / demande (notifications SMS possibles).")
    add_bullet(
        doc,
        "Factures déjà soldées : saisir pour le registre ; ne pas les compter en dette ouverte "
        "sans consigne métier claire.",
    )
    add_bullet(doc, "Contrôle qualité lundi : échantillon minimum 1 courrier sur 10.")

    # 12. Résultat attendu
    add_heading_custom(doc, "12. Résultat attendu en fin de lundi", 1)
    add_bullet(doc, "209 courriers d’août présents dans le registre d’arrivée COSUD.")
    add_bullet(doc, "Factures liées au référentiel fournisseurs / prestataires.")
    add_bullet(doc, "Checklist Excel à jour (statut « saisi » / « contrôlé »).")
    add_bullet(doc, "Liste des éventuels restes / bloquants pour J+1 (scans manquants, etc.).")
    add_bullet(doc, "Base prête pour analyse dettes puis moratoires.")

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    run = footer.add_run(
        "Document COSUD / ACSI — Programme opérationnel de reprise registre août — Commission DG."
    )
    set_run_font(run, size=9, italic=True, color=GRAY)

    doc.save(OUTPUT)
    print(f"OK: {OUTPUT}")


if __name__ == "__main__":
    build()
