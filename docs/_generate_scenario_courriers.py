"""Génère docs/Scenario-test-circuits-courriers.docx depuis docs/_scenario_text.txt."""

from __future__ import annotations

import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = os.path.dirname(__file__)
SOURCE = os.path.join(DOCS_DIR, "_scenario_text.txt")
OUTPUT = os.path.join(DOCS_DIR, "Scenario-test-circuits-courriers.docx")

GREEN = RGBColor(0x00, 0xA0, 0x55)
DARK = RGBColor(0x0F, 0x17, 0x2A)


def is_table_row(line: str) -> bool:
    return " | " in line and not line.strip().startswith("Champ")


def parse_table_block(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = [ln.split(" | ") for ln in lines if " | " in ln]
    if not rows:
        return [], []
    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []
    return header, body


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    if not header:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for col, text in enumerate(header):
        cell = table.rows[0].cells[col]
        cell.text = text.strip()
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            if c_idx < len(header):
                table.rows[r_idx + 1].cells[c_idx].text = text.strip()
    doc.add_paragraph()


def build_document() -> Document:
    with open(SOURCE, encoding="utf-8") as f:
        raw_lines = [ln.rstrip() for ln in f.readlines()]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Scénario de test — Circuits de traitement des courriers")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = GREEN

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Application GED — Module Courriers")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = DARK

    doc.add_paragraph()

    i = 0
    while i < len(raw_lines):
        line = raw_lines[i].strip()
        if not line:
            i += 1
            continue

        # Skip duplicate title lines already rendered
        if line.startswith("Scénario de test") or line == "Application GED — Module Courriers":
            i += 1
            continue

        if re.match(r"^\d+\.\s", line) and not re.match(r"^\d+\.\d+", line):
            p = doc.add_heading(line, level=1)
            for run in p.runs:
                run.font.color.rgb = GREEN
            i += 1
            continue

        if re.match(r"^\d+\.\d+", line):
            doc.add_heading(line, level=2)
            i += 1
            continue

        if is_table_row(line):
            block: list[str] = []
            while i < len(raw_lines) and is_table_row(raw_lines[i].strip()):
                block.append(raw_lines[i].strip())
                i += 1
            header, rows = parse_table_block(block)
            add_table(doc, header, rows)
            continue

        if line.startswith("Attendu :") or line.startswith("Commande :"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line)
            run.bold = True
            i += 1
            continue

        if line.startswith("Important :") or line.startswith("Note :") or line.startswith("Les cas"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            i += 1
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            i += 1
            continue

        doc.add_paragraph(line)
        i += 1

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"OK {OUTPUT}")


if __name__ == "__main__":
    main()
