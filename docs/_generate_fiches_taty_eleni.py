"""Génère docs/Fiches-techniques-Taty-Eleni.docx — fiches métier (validation)."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = os.path.dirname(__file__)
OUTPUT = os.path.join(DOCS_DIR, "Fiches-techniques-Taty-Eleni.docx")
OUTPUT_ALT = os.path.join(DOCS_DIR, "Fiches-techniques-Taty-Eleni-option-A.docx")

GREEN = RGBColor(0x00, 0xA0, 0x55)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)


def set_run_font(run, *, size=11, bold=False, color=DARK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_heading_custom(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=GREEN)
    else:
        set_run_font(run, size=13 if level == 2 else 11, bold=True, color=DARK)
    return p


def add_para(doc: Document, text: str, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    # Clear default run then style
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0])
    else:
        set_run_font(p.add_run(text))


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0])
    else:
        set_run_font(p.add_run(text))


def set_cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_run_font(cell.paragraphs[0].add_run(h), size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "00A055")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            set_run_font(cell.paragraphs[0].add_run(val), size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0FDF4")
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.85)
        s.bottom_margin = Inches(0.85)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    set_run_font(title.add_run("FICHES TECHNIQUES MÉTIER"), size=20, bold=True, color=GREEN)

    subt = doc.add_paragraph()
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        subt.add_run("Secrétariat général — Dossiers fournisseurs (Taty) et suivi des dépenses (Eleni)"),
        size=13,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(12)
    set_run_font(
        meta.add_run(
            "Agence Congolaise des Systèmes d’Information (ACSI)\n"
            "Document de validation métier — août 2026\n"
            "Version : 1.1 — périmètre Eleni élargi ; classement facture = Taty seule ; MAD hors classement dossier"
        ),
        size=10,
        color=GRAY,
        italic=True,
    )

    # --- Principes ---
    add_heading_custom(doc, "Principes communs", 1)
    add_bullet(
        doc,
        "Tous les secrétaires / acteurs secrétariat peuvent enregistrer un courrier arrivée "
        "(facture, MAD, demande, etc.).",
    )
    add_bullet(
        doc,
        "Seule la responsable de la tâche agit ensuite : classement dossier (Taty) ou "
        "contrôle / clôture et suivi financier (Eleni).",
    )
    add_bullet(
        doc,
        "Sur facture ou MAD : le DG instruit uniquement via le panneau « Circuit métier » "
        "(Bon pour accord). Pas de « Répondre moi-même » qui clôturerait le circuit paiement.",
    )
    add_bullet(
        doc,
        "Le N° fulgurant n’est plus saisi à l’enregistrement (recherche legacy éventuelle uniquement).",
    )

    add_heading_custom(doc, "Enchaînement après Bon pour accord (facture / MAD)", 1)
    add_numbered(doc, "Le DG donne son Bon pour accord (instructions ± confier à un directeur).")
    add_numbered(doc, "L’Agent comptable établit le chèque et l’envoie au DG (papier).")
    add_numbered(
        doc,
        "Le DG signe le chèque sur papier, puis confirme dans l’outil "
        "« Chèque signé — renvoyer à l’AC » (sans scan).",
    )
    add_numbered(
        doc,
        "Lorsque le bénéficiaire décharge le chèque, l’AC enregistre le bordereau "
        "(date, n° pièce, montant, banque, bénéficiaire, programmation) et joint les pièces "
        "(chèque déchargé, pièce d’identité, etc.).",
    )
    add_numbered(
        doc,
        "Mme Eleni Ossebi est notifiée : elle contrôle les éléments saisis avec les pièces "
        "physiques, peut ajouter des pièces complémentaires, puis confirme la clôture.",
    )
    add_para(
        doc,
        "En parallèle dès le Bon pour accord (uniquement pour une FACTURE) : Mme Taty "
        "classe la facture dans le dossier fournisseur et suit le paiement — elle n’envoie "
        "pas le dossier à l’AC. Une MAD n’ouvre pas l’action « Classer dans un dossier ».",
        italic=True,
    )

    # --- Taty ---
    doc.add_page_break()
    add_heading_custom(doc, "FICHE 1 — Mme ANNE LETHICIA TATY-TCHICAYA", 1)
    add_para(doc, "Rôle GED : responsable_dossiers_prestataires", bold=True)
    add_para(doc, "Responsable des dossiers fournisseurs / prestataires", bold=True)

    add_heading_custom(doc, "1. Objectif", 2)
    add_para(
        doc,
        "Classer les factures approuvées dans le dossier GED du fournisseur, suivre "
        "ponctuellement les paiements factures, et produire un rapport chaque vendredi "
        "à la demande du DG (périmètre factures fournisseurs).",
    )

    add_heading_custom(doc, "2. Périmètre", 2)
    add_bullet(doc, "Factures prestataires / fournisseurs uniquement.")
    add_bullet(doc, "Hors périmètre classement : MAD, demandes générales, autres courriers.")
    add_bullet(
        doc,
        "Peut enregistrer des courriers comme les autres acteurs secrétariat ; "
        "le classement dossier reste sa responsabilité exclusive (hors admin).",
    )

    add_heading_custom(doc, "3. Déclencheur", 2)
    add_para(doc, "Bon pour accord du DG sur une facture (notification / suivi parallèle).")

    add_heading_custom(doc, "4. Actions dans le GED", 2)
    add_numbered(doc, "Prendre connaissance de la facture approuvée.")
    add_numbered(
        doc,
        "Classer la facture (et ses pièces) dans un dossier fournisseur existant "
        "ou créer le dossier sous « Mes dossiers ».",
    )
    add_numbered(doc, "Suivre l’avancement du paiement (page Factures fournisseurs).")
    add_numbered(doc, "Chaque vendredi : exporter / produire le rapport de suivi pour le DG.")

    add_heading_custom(doc, "5. Ce qu’elle ne fait pas", 2)
    add_bullet(doc, "Elle n’envoie pas le dossier à l’Agent comptable.")
    add_bullet(doc, "Elle ne bloque pas le paiement.")
    add_bullet(doc, "Elle n’enregistre pas la décharge ni ne clôture la dépense (rôle Eleni / AC).")
    add_bullet(doc, "Elle ne classe pas une MAD dans un dossier fournisseur.")

    # --- Eleni ---
    doc.add_page_break()
    add_heading_custom(doc, "FICHE 2 — Mme ASTRIDE ELENI OSSEBI", 1)
    add_para(doc, "Rôle GED : responsable_suivi_depenses", bold=True)
    add_para(doc, "Responsable du suivi des dépenses — périmètre finances", bold=True)

    add_heading_custom(doc, "1. Objectif", 2)
    add_para(
        doc,
        "Assurer le suivi de tout ce qui concerne les finances placées sous sa "
        "responsabilité : contrôler les dépenses enregistrées, tenir les fiches de suivi, "
        "et produire chaque semaine un rapport selon la demande du DG.",
    )

    add_heading_custom(doc, "2. Périmètre métier (validation Eleni — août 2026)", 2)
    add_para(doc, "Au-delà du suivi des factures / MAD déjà dans le circuit paiement :", italic=True)
    add_bullet(
        doc,
        "Paiement des éléments liés à la paie — exemples : pharmacie AFCOM, "
        "remboursement de frais médicaux, achat de lunettes et assimilés.",
    )
    add_bullet(doc, "Les commissions.")
    add_bullet(doc, "La TTF.")
    add_bullet(
        doc,
        "Toute autre dépense dont le DG lui remet les copies pour suivi financier.",
    )
    add_para(
        doc,
        "Canal physique : à la sortie du bureau du DG, les copies lui sont remises "
        "personnellement, sans intermédiaire, pour qu’elle suive ce genre de dépenses.",
        italic=True,
    )

    add_heading_custom(doc, "3. Déclencheurs", 2)
    add_bullet(
        doc,
        "Circuit facture / MAD : notification après que l’AC a enregistré la décharge "
        "(bordereau + pièces) — contrôle puis clôture.",
    )
    add_bullet(
        doc,
        "Autres dépenses (paie, commissions, TTF…) : remise directe des copies par le DG "
        "(à formaliser dans le GED : types / fiches / dépôt dédié — à implémenter).",
    )

    add_heading_custom(doc, "4. Actions — circuit facture / MAD (déjà en place)", 2)
    add_numbered(
        doc,
        "Consulter le bordereau saisi par l’AC (date, n° pièce, montant, banque, "
        "bénéficiaire, programmation).",
    )
    add_numbered(doc, "Comparer avec les pièces physiques en sa possession.")
    add_numbered(doc, "Joindre éventuellement des pièces complémentaires.")
    add_numbered(doc, "Confirmer le contrôle → clôture du dossier.")
    add_numbered(doc, "Consulter / exporter les fiches de suivi des paiements (FSP facture, FSP MAD).")

    add_heading_custom(doc, "5. Actions — hors factures (à couvrir dans le GED)", 2)
    add_numbered(doc, "Enregistrer / rattacher les copies remises par le DG (paie, commissions, TTF…).")
    add_numbered(doc, "Suivre l’état de chaque dépense jusqu’à clôture / contrôle.")
    add_numbered(
        doc,
        "Chaque semaine : rapport consolidé au DG (toutes les lignes de son périmètre "
        "finances, pas seulement les factures fournisseurs de Taty).",
    )

    add_heading_custom(doc, "6. Ce qu’elle ne fait pas", 2)
    add_bullet(doc, "Elle n’enregistre pas la décharge à la place de l’AC (circuit facture/MAD).")
    add_bullet(doc, "Elle ne signe pas les chèques.")
    add_bullet(doc, "Le DG ne lui envoie pas de scan de chèque signé (confirmation GED sans scan).")
    add_bullet(doc, "Elle ne classe pas les factures dans les dossiers fournisseurs (rôle Taty).")

    # --- AC ---
    doc.add_page_break()
    add_heading_custom(doc, "Rôle de l’Agent comptable (rappel)", 1)
    add_table(
        doc,
        ["Moment", "Action"],
        [
            ["Après Bon pour accord", "Établit le chèque → envoi au DG"],
            ["Après signature DG", "Attend la décharge du bénéficiaire"],
            [
                "À la décharge",
                "Saisit le bordereau + joint chèque déchargé, identité, etc. "
                "(enregistrement de la preuve de paiement)",
            ],
        ],
    )

    add_heading_custom(doc, "Bordereau de transmission (données à saisir)", 2)
    add_table(
        doc,
        ["Colonne", "Exemple"],
        [
            ["Date", "21/07/2026"],
            ["N° pièce", "Chèque N° 0000312 / Mise à disposition de fonds"],
            ["Montant", "1 000 000"],
            ["Banque", "BCH, BOA…"],
            ["Bénéficiaire", "BL Technology / responsable de département"],
            ["Programmation", "du 14 juillet 2026"],
            ["Émargement", "Pièces scannées (décharge + identité…)"],
        ],
    )

    add_heading_custom(doc, "Synthèse des responsabilités", 1)
    add_table(
        doc,
        ["", "Mme Taty", "Agent comptable", "Mme Eleni"],
        [
            [
                "Périmètre",
                "Dossiers + suivi factures fournisseurs",
                "Chèque + décharge (circuit paiement)",
                "Finances : factures/MAD + paie + commissions + TTF + autres dépenses DG",
            ],
            [
                "Déclencheur",
                "BPA DG sur facture",
                "BPA puis chèque signé",
                "Décharge AC (circuit) ou copies remises par le DG",
            ],
            [
                "Action clé",
                "Classer dossier + rapport vendredi factures",
                "Bordereau + pièces à la décharge",
                "Contrôle / clôture + FSP + rapport hebdo finances",
            ],
            [
                "MAD",
                "Pas de classement dossier fournisseur",
                "Comme facture si circuit paiement",
                "Suivi / contrôle (et FSP MAD)",
            ],
            ["Scan chèque signé DG", "—", "Non (à la décharge seulement)", "Contrôle des pièces AC"],
        ],
    )

    add_heading_custom(doc, "Évolutions GED prévues (suite à validation)", 1)
    add_bullet(
        doc,
        "Formaliser dans l’outil les catégories hors facture (éléments de paie, commissions, TTF).",
    )
    add_bullet(
        doc,
        "Canal DG → Eleni (remise des copies) traçable sans intermédiaire obligatoire.",
    )
    add_bullet(
        doc,
        "Rapport hebdomadaire consolidé du périmètre Eleni (au-delà de la page Factures fournisseurs).",
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    set_run_font(
        note.add_run(
            "Document établi pour validation métier — Secrétariat général / ACSI — août 2026 "
            "(v1.1). À valider avec Mme Taty et Mme Eleni avant poursuite d’implémentation."
        ),
        size=9,
        italic=True,
        color=GRAY,
    )

    try:
        doc.save(OUTPUT)
        print(f"OK: {OUTPUT}")
    except PermissionError:
        doc.save(OUTPUT_ALT)
        print(f"OK (alt): {OUTPUT_ALT}")


if __name__ == "__main__":
    build()
