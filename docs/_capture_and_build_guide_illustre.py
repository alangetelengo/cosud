"""
Capture les écrans COSUD (http://cosud) et génère un guide Word illustré.
Usage: python docs/_capture_and_build_guide_illustre.py
"""

from __future__ import annotations

import os
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import TimeoutError as PwTimeout
from playwright.sync_api import sync_playwright

BASE = Path(__file__).resolve().parent
SHOTS = BASE / "captures-guide"
OUT_DOC = BASE / "Guide-utilisateur-illustre-COSUD.docx"
BASE_URL = os.environ.get("COSUD_URL", "http://cosud").rstrip("/")

GREEN = RGBColor(0x00, 0xA0, 0x55)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)

# Comptes démo (mot de passe commun seed)
PASSWORD = "password"

# (slug fichier, titre section, email login, chemin relatif, note)
# Comptes démo seed UsersAvecDirectionsSeeder — mot de passe « password »
SCREENS: list[tuple[str, str, str, str, str]] = [
    ("01-login", "Connexion", "", "/login", "Page d’accès COSUD."),
    (
        "02-dashboard-admin",
        "Tableau de bord (Administrateur)",
        "admin@acsi.cg",
        "/",
        "Vue d’accueil après connexion — menus complets.",
    ),
    (
        "03-documents",
        "Documents",
        "admin@acsi.cg",
        "/documents",
        "GED — liste des documents.",
    ),
    (
        "04-dossiers",
        "Dossiers",
        "admin@acsi.cg",
        "/dossiers",
        "Arborescence / liste des dossiers.",
    ),
    (
        "05-courriers",
        "Courriers — liste",
        "admin@acsi.cg",
        "/courriers",
        "Liste des courriers.",
    ),
    (
        "06-courrier-create",
        "Courriers — nouveau (arrivée)",
        "admin@acsi.cg",
        "/courriers/create?sens=arrivee",
        "Formulaire d’enregistrement d’un courrier arrivée (scan obligatoire).",
    ),
    (
        "07-registre-arrivee",
        "Registre Arrivée",
        "admin@acsi.cg",
        "/registres/courriers/arrivee",
        "Registre officiel des arrivées.",
    ),
    (
        "08-registre-depart",
        "Registre Départ",
        "admin@acsi.cg",
        "/registres/courriers/depart",
        "Registre officiel des départs.",
    ),
    (
        "09-fournisseurs",
        "Fournisseurs / prestataires",
        "admin@acsi.cg",
        "/fournisseurs-prestataires",
        "Référentiel obligatoire pour les factures.",
    ),
    (
        "10-fournisseurs-create",
        "Créer une fiche fournisseur",
        "admin@acsi.cg",
        "/fournisseurs-prestataires/create",
        "Création d’une fiche référentiel.",
    ),
    (
        "11-suivi-factures",
        "Suivi de factures",
        "admin@acsi.cg",
        "/suivi-factures-fournisseurs",
        "Suivi des factures fournisseurs / dettes.",
    ),
    (
        "12-suivi-depenses",
        "Suivi de dépense",
        "admin@acsi.cg",
        "/suivi-paiements",
        "Suivi des paiements / dépenses.",
    ),
    (
        "13-regularisation",
        "Reprise des factures prestataires",
        "admin@acsi.cg",
        "/factures-regularisation",
        "Module hors circuit (historique / dette).",
    ),
    (
        "14-moratoires",
        "Moratoires",
        "admin@acsi.cg",
        "/moratoires",
        "Plans de moratoire fournisseurs.",
    ),
    (
        "15-bordereau",
        "Bordereau de transmission",
        "admin@acsi.cg",
        "/bordereau-transmission",
        "Bordereaux de transmission.",
    ),
    (
        "16-parametres",
        "Paramètres",
        "admin@acsi.cg",
        "/parametres",
        "Centre d’administration.",
    ),
    (
        "17-notifications",
        "Paramètres — Notifications",
        "admin@acsi.cg",
        "/parametres/notifications",
        "Activer / désactiver la notif DG (SMS + cloche) à l’enregistrement facture/MAD.",
    ),
    (
        "18-utilisateurs",
        "Utilisateurs",
        "admin@acsi.cg",
        "/utilisateurs",
        "Gestion des comptes.",
    ),
    (
        "19-sidebar-secretaire",
        "Menus — Secrétaire de direction",
        "chef.sec@acsi.cg",
        "/",
        "Exemple de menus pour le rôle secrétaire (courriers + registres, sans admin technique).",
    ),
]


def set_run_font(run, *, size=11, bold=False, color=DARK, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def login(page, email: str) -> None:
    page.goto(f"{BASE_URL}/login", wait_until="domcontentloaded", timeout=60000)
    page.fill("#email", email)
    page.fill("#password", PASSWORD)
    with page.expect_navigation(wait_until="networkidle", timeout=60000):
        page.click('button[type="submit"]')
    if "/login" in page.url:
        # Afficher erreur éventuelle pour debug
        err = ""
        try:
            err = page.locator(".err").inner_text(timeout=1000)
        except Exception:
            pass
        raise RuntimeError(f"Echec connexion pour {email}: {err or page.url}")


def logout(page) -> None:
    # Tentative logout via lien / formulaire
    try:
        page.goto(f"{BASE_URL}/logout", wait_until="domcontentloaded", timeout=15000)
    except Exception:
        pass
    # Breeze utilise souvent POST logout — essayer formulaire si présent
    try:
        page.goto(f"{BASE_URL}/login", wait_until="domcontentloaded", timeout=30000)
    except Exception:
        pass


def capture_all() -> list[dict]:
    SHOTS.mkdir(parents=True, exist_ok=True)
    results: list[dict] = []

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            viewport={"width": 1440, "height": 900},
            locale="fr-FR",
        )
        page = context.new_page()
        current_user: str | None = None

        for slug, title, email, path, note in SCREENS:
            out = SHOTS / f"{slug}.png"
            ok = False
            err = None
            try:
                if not email:
                    page.goto(f"{BASE_URL}{path}", wait_until="networkidle", timeout=60000)
                else:
                    if current_user != email:
                        context.clear_cookies()
                        login(page, email)
                        current_user = email
                    page.goto(f"{BASE_URL}{path}", wait_until="networkidle", timeout=60000)
                    if "/login" in page.url:
                        # Session perdue : reconnecter une fois
                        context.clear_cookies()
                        login(page, email)
                        current_user = email
                        page.goto(f"{BASE_URL}{path}", wait_until="networkidle", timeout=60000)
                    if "/login" in page.url:
                        raise RuntimeError(f"Redirection login sur {path}")
                time.sleep(0.7)
                page.screenshot(path=str(out), full_page=False)
                ok = out.exists() and out.stat().st_size > 1000
            except Exception as e:
                err = str(e)
                try:
                    page.screenshot(path=str(out), full_page=False)
                    ok = out.exists()
                except Exception:
                    ok = False

            results.append(
                {
                    "slug": slug,
                    "title": title,
                    "email": email or "(public)",
                    "path": path,
                    "note": note,
                    "file": out if ok else None,
                    "error": err,
                }
            )
            print(("OK" if ok else "KO"), slug, err or "")

        browser.close()
    return results


def build_doc(results: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("COSUD — ACSI")
    set_run_font(r, size=14, bold=True, color=GREEN)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Guide utilisateur illustré")
    set_run_font(r2, size=22, bold=True, color=DARK)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(
        "Captures d’écran des principaux niveaux de l’application\n"
        f"Base : {BASE_URL} — Guide permanent"
    )
    set_run_font(r3, size=10, italic=True, color=GRAY)

    intro = doc.add_paragraph()
    ri = intro.add_run(
        "Ce document présente les écrans clés de COSUD. "
        "Les captures sont organisées par niveau fonctionnel. "
        "Aucun nom de responsable n’est cité : seuls les rôles / menus sont décrits."
    )
    set_run_font(ri, size=11)

    som = doc.add_paragraph()
    rs = som.add_run("Sommaire des écrans")
    set_run_font(rs, size=14, bold=True, color=GREEN)
    for i, item in enumerate(results, start=1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item["title"])
        set_run_font(run, size=11)

    for i, item in enumerate(results, start=1):
        doc.add_page_break()
        h = doc.add_paragraph()
        rh = h.add_run(f"{i}. {item['title']}")
        set_run_font(rh, size=16, bold=True, color=GREEN)

        meta = doc.add_paragraph()
        rm = meta.add_run(f"URL : {item['path']}    |    Compte de capture : {item['email']}")
        set_run_font(rm, size=9, italic=True, color=GRAY)

        note = doc.add_paragraph()
        rn = note.add_run(item["note"])
        set_run_font(rn, size=11)

        if item["file"] and Path(item["file"]).exists():
            doc.add_picture(str(item["file"]), width=Inches(6.5))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            errp = doc.add_paragraph()
            re = errp.add_run(
                "Capture indisponible"
                + (f" ({item['error']})" if item.get("error") else "")
                + ". Ouvrir l’écran manuellement dans COSUD."
            )
            set_run_font(re, size=11, italic=True, color=GRAY)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    rf = footer.add_run(
        "Document COSUD / ACSI — Guide utilisateur illustré — généré automatiquement."
    )
    set_run_font(rf, size=9, italic=True, color=GRAY)

    doc.save(str(OUT_DOC))
    print(f"DOC: {OUT_DOC}")


def main() -> None:
    print(f"BASE_URL={BASE_URL}")
    results = capture_all()
    build_doc(results)
    ok_n = sum(1 for r in results if r["file"])
    print(f"Captures OK: {ok_n}/{len(results)}")


if __name__ == "__main__":
    main()
